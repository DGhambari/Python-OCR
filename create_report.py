from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
from datetime import datetime
import read_sequence
import compare_files
import os.path

hmi_list = './json/HMI_valve_list.json'
doc_list = './json/doc_valve_list.json'

hmi_valves = compare_files.retrieve_valve_list(hmi_list)
doc_valves = compare_files.retrieve_valve_list(doc_list)
step_number = read_sequence.read_step_num()
table_length = len(hmi_valves)

img_path = f'./images/Step_{step_number}.PNG'
report_path = './reports'

# Create a new Word document
doc = Document()

# Add a title and a heading
doc.add_heading(f'Step {step_number} Valve State Verification', level=1)
doc.add_heading('Valve State Screenshot', level=2)

# Add screenshot of the system
doc.add_picture(img_path, width=Cm(16))
doc.add_paragraph('')

# Move onto the next page
doc.add_page_break()

# Create a table and add two rows to allow for headings
doc.add_heading('Valve State Comparison Results', level=2)
table = doc.add_table(rows=table_length+3, cols=4)
table.style = 'Table Grid'

# Set the title and header rows
title_row = table.rows[0]
title_cell = title_row.cells[0]
title_cell.text = f'Step {step_number} Results'
title_cell.paragraphs[0].alignment = 1

header_row = table.rows[1]
header_row.cells[0].text = 'HMI / Live System Information'
header_row.cells[2].text = 'Design Specification'

header_row_2 = table.rows[2]
header_row_2.cells[0].text = 'Valve Tag'
header_row_2.cells[1].text = 'Valve State'
header_row_2.cells[2].text = 'Valve Tag'
header_row_2.cells[3].text = 'Valve State'

# Merge the title cell across all columns
for cell in title_row.cells:
    cell.width = Cm(10)
title_row.cells[0].merge(title_row.cells[3])

# Merge the header cells across 2 columns
header_row.cells[0].merge(header_row.cells[1])
header_row.cells[2].merge(header_row.cells[3])

# Populate the table by iterating through HMI valve states
row_index = 2
for valve, hmi_state in hmi_valves.items():
    row_index += 1
    row = table.rows[row_index].cells
    row[0].text = valve
    row[1].text = hmi_state

    if valve in doc_valves:
        row[2].text = valve
        row[3].text = doc_valves[valve]

# Populate the table by iterating through the doc_valve list
row_index = 2
for valve, doc_state in doc_valves.items():
    if valve not in hmi_valves:
        row_index += 1
        row = table.rows[row_index].cells
        if not row[2].text:
            row[2].text = valve
            row[3].text = doc_state
        else:
            row[2].text = ""
            row[3].text = ""

column_width = Cm(2.5)  

# Set the width for each column
for col in table.columns:
    for cell in col.cells:
        cell.width = column_width

# Get current date and time for the report filename
current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

# Save the document with the date and time in the filename
report_filename = f'Step {step_number} Report - {current_datetime}.docx'
if os.path.exists(report_path) == False:
    os.mkdir(report_path)
doc.save(f'{report_path}/{report_filename}')

print(f"Report saved as '{report_filename}'")
